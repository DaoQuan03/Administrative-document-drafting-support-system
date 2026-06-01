import json
import os
import random
import requests
import urllib.parse
import sys
from google import genai
from google.genai import errors as google_errors
from bs4 import BeautifulSoup
from pathlib import Path

# Thiết lập encoding UTF-8 cho console để hiển thị mượt mà trên Windows
if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except AttributeError:
        pass
from django.shortcuts import render, redirect
from django.http import JsonResponse, FileResponse
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth import get_user_model
from django.views.decorators.csrf import csrf_exempt, ensure_csrf_cookie
from django.contrib.auth.decorators import login_required
from django.core.mail import send_mail
from django.conf import settings

from .models import Document, KnowledgeBase, KnowledgeDocument

User = get_user_model()

@ensure_csrf_cookie
def serve_html_page(request, page_name=None):
    # If root path, redirect depending on auth status
    if not page_name:
        if request.user.is_authenticated:
            return redirect('/html/dashboard.html')
        else:
            return redirect('/html/login.html')
            
    # Normalize path
    if not page_name.endswith('.html'):
        page_name += '.html'
        
    # Login and register pages are accessible to everyone
    if page_name in ['login.html', 'register.html']:
        if request.user.is_authenticated and page_name == 'login.html':
            return redirect('/html/dashboard.html')
        return render(request, page_name)
        
    # All other dashboard / editor / settings pages require authentication
    if not request.user.is_authenticated:
        return redirect('/html/login.html')
        
    # Serve requested template
    try:
        return render(request, page_name)
    except Exception:
        # Fallback if page not found
        return redirect('/html/dashboard.html')

@csrf_exempt
def api_login(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        username = data.get('email')  # Can be username or email
        password = data.get('password')
    except Exception:
        return JsonResponse({'success': False, 'error': 'Dữ liệu JSON không hợp lệ.'}, status=400)
        
    if not username or not password:
        return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp đầy đủ thông tin.'}, status=400)
        
    user = authenticate(request, username=username, password=password)
    if user is not None:
        login(request, user)
        return JsonResponse({
            'success': True,
            'user': {
                'username': user.username,
                'email': user.email,
                'full_name': user.full_name,
                'role': user.role,
                'organization': user.organization
            }
        })
    else:
        return JsonResponse({'success': False, 'error': 'Tên đăng nhập/Email hoặc mật khẩu không chính xác.'}, status=400)

@csrf_exempt
def api_register(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        full_name = data.get('full_name')
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')
        role = data.get('role', 'MEMBER')
        organization = data.get('organization', 'Cá nhân')
        title = data.get('title', '')
        department = data.get('department', '')
    except Exception:
        return JsonResponse({'success': False, 'error': 'Dữ liệu JSON không hợp lệ.'}, status=400)
        
    # Basic validation
    if not full_name or not username or not email or not password:
        return JsonResponse({'success': False, 'error': 'Vui lòng điền đầy đủ các thông tin bắt buộc.'}, status=400)
        
    # Check if username or email already exists
    if User.objects.filter(username__iexact=username).exists():
        return JsonResponse({'success': False, 'error': 'Tên đăng nhập này đã tồn tại.'}, status=400)
        
    if User.objects.filter(email__iexact=email).exists():
        return JsonResponse({'success': False, 'error': 'Email này đã được đăng ký.'}, status=400)
        
    # Create user
    try:
        user = User.objects.create_user(
            username=username,
            email=email,
            password=password,
            full_name=full_name,
            role=role,
            organization=organization or 'Cá nhân',
            title=title,
            department=department
        )
        # Do not automatically log in. Return success message for redirect.
        return JsonResponse({
            'success': True,
            'message': 'Đăng ký tài khoản thành công! Vui lòng đăng nhập với tài khoản của bạn.'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)

@csrf_exempt
def api_logout(request):
    logout(request)
    if request.method == 'GET':
        return redirect('/html/login.html')
    return JsonResponse({'success': True})

@csrf_exempt
def api_update_profile(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        full_name = data.get('full_name')
        title = data.get('title', '')
        organization = data.get('organization', '')
        department = data.get('department', '')
        phone_number = data.get('phone_number', '').strip()
        birth_date_str = data.get('birth_date', '').strip()
        confirm_password = data.get('confirm_password', '')
    except Exception:
        return JsonResponse({'success': False, 'error': 'Dữ liệu JSON không hợp lệ.'}, status=400)
        
    if not full_name:
        return JsonResponse({'success': False, 'error': 'Họ và tên không được để trống.'}, status=400)
        
    birth_date = None
    if birth_date_str:
        from datetime import datetime
        try:
            birth_date = datetime.strptime(birth_date_str, '%Y-%m-%d').date()
        except ValueError:
            return JsonResponse({'success': False, 'error': 'Ngày sinh không đúng định dạng YYYY-MM-DD.'}, status=400)
            
    user = request.user
    
    # Yêu cầu xác nhận mật khẩu (nếu tài khoản có mật khẩu thực tế)
    if user.has_usable_password():
        if not confirm_password:
            return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp mật khẩu hiện tại để xác nhận thay đổi.'}, status=400)
        if not user.check_password(confirm_password):
            return JsonResponse({'success': False, 'error': 'Mật khẩu xác nhận không chính xác.'}, status=400)
            
    try:
        user.full_name = full_name
        user.title = title
        if organization:
            user.organization = organization
        user.department = department
        user.phone_number = phone_number
        user.birth_date = birth_date
        user.save()
        
        return JsonResponse({
            'success': True,
            'user': {
                'username': user.username,
                'email': user.email,
                'full_name': user.full_name,
                'role': user.role,
                'organization': user.organization,
                'title': user.title,
                'department': user.department,
                'phone_number': user.phone_number,
                'birth_date': user.birth_date.strftime('%Y-%m-%d') if user.birth_date else '',
                'avatar': user.avatar
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)


@csrf_exempt
def api_upload_avatar(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        avatar_file = request.FILES.get('avatar')
        if not avatar_file:
            return JsonResponse({'success': False, 'error': 'Không nhận được file ảnh.'}, status=400)
            
        import os
        from pathlib import Path
        
        ext = os.path.splitext(avatar_file.name)[1].lower()
        if ext not in ['.png', '.jpg', '.jpeg', '.gif']:
            return JsonResponse({'success': False, 'error': 'Định dạng ảnh không hợp lệ (Chỉ hỗ trợ PNG, JPG, JPEG, GIF).'}, status=400)
            
        if avatar_file.size > 2 * 1024 * 1024:
            return JsonResponse({'success': False, 'error': 'Kích thước file vượt quá giới hạn 2MB.'}, status=400)
            
        BASE_DIR = Path(__file__).resolve().parent.parent
        upload_dir = os.path.join(BASE_DIR, 'frontend', 'uploads', 'avatars')
        os.makedirs(upload_dir, exist_ok=True)
        
        filename = f"avatar_{request.user.id}{ext}"
        filepath = os.path.join(upload_dir, filename)
        
        # Save file chunks
        with open(filepath, 'wb+') as destination:
            for chunk in avatar_file.chunks():
                destination.write(chunk)
                
        relative_path = f"/static/uploads/avatars/{filename}"
        
        # Update user's avatar path in DB
        user = request.user
        user.avatar = relative_path
        user.save()
        
        return JsonResponse({
            'success': True,
            'avatar_url': relative_path
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)

@csrf_exempt
def api_register_send_code(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        full_name = data.get('full_name')
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')
        role = data.get('role', 'MEMBER')
        organization = data.get('organization', 'Cá nhân')
        title = data.get('title', '')
        department = data.get('department', '')
    except Exception:
        return JsonResponse({'success': False, 'error': 'Dữ liệu JSON không hợp lệ.'}, status=400)
        
    if not full_name or not username or not email or not password:
        return JsonResponse({'success': False, 'error': 'Vui lòng điền đầy đủ các thông tin bắt buộc.'}, status=400)
        
    if User.objects.filter(username__iexact=username).exists():
        return JsonResponse({'success': False, 'error': 'Tên đăng nhập này đã tồn tại.'}, status=400)
        
    if User.objects.filter(email__iexact=email).exists():
        return JsonResponse({'success': False, 'error': 'Email này đã được đăng ký.'}, status=400)
        
    # Generate 6-digit OTP
    otp_code = str(random.randint(100000, 999999))
    
    # Store registration details & OTP in session
    request.session['pending_register'] = {
        'full_name': full_name,
        'username': username,
        'email': email,
        'password': password,
        'role': role,
        'organization': organization,
        'title': title,
        'department': department,
        'otp': otp_code
    }
    request.session.set_expiry(600)  # Expire in 10 minutes
    
    subject = "[VănBảnAI] Mã xác minh đăng ký tài khoản mới"
    message = f"Xin chào {full_name},\n\nMã OTP xác minh đăng ký tài khoản của bạn tại hệ thống VănBảnAI là: {otp_code}\n\nMã xác minh này có hiệu lực trong vòng 10 phút. Vui lòng không chia sẻ mã này với bất kỳ ai.\n\nTrân trọng,\nĐội ngũ VănBảnAI."
    
    try:
        send_mail(
            subject,
            message,
            settings.DEFAULT_FROM_EMAIL,
            [email],
            fail_silently=False,
        )
        return JsonResponse({
            'success': True, 
            'message': 'Mã OTP đã được gửi đến email của bạn.',
            'otp_local_dev': otp_code
        })
    except Exception as e:
        # Fallback logging
        print(f"==================================================")
        print(f"SMTP FAILED: {str(e)}")
        print(f"OTP CODE FOR EMAIL {email}: {otp_code}")
        print(f"==================================================")
        return JsonResponse({
            'success': True,
            'message': 'Mã OTP đã được cấp phát. Do đang ở môi trường phát triển local, vui lòng xem mã OTP trong Terminal chạy Server hoặc Hộp thư của bạn.',
            'otp_local_dev': otp_code
        })

@csrf_exempt
def api_register_verify(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    pending_data = request.session.get('pending_register')
    if not pending_data:
        return JsonResponse({'success': False, 'error': 'Không tìm thấy thông tin đăng ký tạm thời hoặc mã OTP đã hết hạn.'}, status=400)
        
    try:
        data = json.loads(request.body)
        otp_entered = data.get('otp')
    except Exception:
        return JsonResponse({'success': False, 'error': 'Dữ liệu JSON không hợp lệ.'}, status=400)
        
    if not otp_entered:
        return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp mã OTP 6 số.'}, status=400)
        
    if str(otp_entered).strip() != str(pending_data['otp']).strip():
        return JsonResponse({'success': False, 'error': 'Mã OTP không chính xác. Vui lòng kiểm tra lại.'}, status=400)
        
    if User.objects.filter(username__iexact=pending_data['username']).exists():
        return JsonResponse({'success': False, 'error': 'Tên đăng nhập này đã tồn tại.'}, status=400)
        
    if User.objects.filter(email__iexact=pending_data['email']).exists():
        return JsonResponse({'success': False, 'error': 'Email này đã được đăng ký.'}, status=400)
        
    try:
        user = User.objects.create_user(
            username=pending_data['username'],
            email=pending_data['email'],
            password=pending_data['password'],
            full_name=pending_data['full_name'],
            role=pending_data['role'],
            organization=pending_data['organization'] or 'Cá nhân',
            title=pending_data['title'],
            department=pending_data['department']
        )
        
        # Clear session
        del request.session['pending_register']
        
        # Do not automatically log in. Return success message for redirect.
        return JsonResponse({
            'success': True,
            'message': 'Đăng ký tài khoản thành công! Vui lòng đăng nhập với tài khoản của bạn.'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)

def google_login(request):
    params = {
        'client_id': settings.GOOGLE_CLIENT_ID,
        'redirect_uri': settings.GOOGLE_REDIRECT_URI,
        'response_type': 'code',
        'scope': 'openid email profile',
        'access_type': 'offline',
        'prompt': 'consent'
    }
    url = "https://accounts.google.com/o/oauth2/v2/auth?" + urllib.parse.urlencode(params)
    return redirect(url)

@csrf_exempt
def google_callback(request):
    code = request.GET.get('code')
    error = request.GET.get('error')
    if error or not code:
        err_msg = f'Lỗi xác thực Google: {error or "Không có mã xác thực"}'
        return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
    
    token_url = "https://oauth2.googleapis.com/token"
    data = {
        'code': code,
        'client_id': settings.GOOGLE_CLIENT_ID,
        'client_secret': settings.GOOGLE_CLIENT_SECRET,
        'redirect_uri': settings.GOOGLE_REDIRECT_URI,
        'grant_type': 'authorization_code'
    }
    
    try:
        response = requests.post(token_url, data=data)
        response_data = response.json()
        print("GOOGLE TOKEN RESPONSE:", response_data)  # Log phản hồi thật của Google
        access_token = response_data.get('access_token')
        if not access_token:
            err_msg = f"Không thể lấy Access Token từ Google. Chi tiết: {response_data.get('error_description', response_data.get('error', 'Không xác định'))}"
            return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
            
        profile_url = "https://www.googleapis.com/oauth2/v3/userinfo"
        headers = {'Authorization': f'Bearer {access_token}'}
        profile_response = requests.get(profile_url, headers=headers)
        profile = profile_response.json()
        
        email = profile.get('email')
        uid = profile.get('sub')
        full_name = profile.get('name', 'Google User')
        
        if not email or not uid:
            err_msg = 'Không lấy được email/ID từ tài khoản Google.'
            return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
            
        try:
            user = User.objects.get(oauth_provider='google', oauth_uid=uid)
            login(request, user, backend='accounts.backends.EmailOrUsernameModelBackend')
            return redirect('/html/dashboard.html')
        except User.DoesNotExist:
            if User.objects.filter(email__iexact=email).exists():
                user = User.objects.get(email__iexact=email)
                if not user.oauth_provider:
                    user.oauth_provider = 'google'
                    user.oauth_uid = uid
                    user.save()
                    login(request, user, backend='accounts.backends.EmailOrUsernameModelBackend')
                    return redirect('/html/dashboard.html')
                else:
                    err_msg = 'Email này đã liên kết với một tài khoản xã hội khác.'
                    return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
            
            request.session['pending_oauth'] = {
                'provider': 'google',
                'uid': uid,
                'email': email,
                'full_name': full_name
            }
            return redirect('/html/login.html?social_register=1')
            
    except Exception as e:
        err_msg = f'Có lỗi xảy ra trong quá trình xác thực Google: {str(e)}'
        return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')

def microsoft_login(request):
    params = {
        'client_id': settings.MICROSOFT_CLIENT_ID,
        'response_type': 'code',
        'redirect_uri': settings.MICROSOFT_REDIRECT_URI,
        'response_mode': 'query',
        'scope': 'https://graph.microsoft.com/User.Read email profile openid',
        'prompt': 'consent'
    }
    url = "https://login.microsoftonline.com/common/oauth2/v2.0/authorize?" + urllib.parse.urlencode(params)
    return redirect(url)

@csrf_exempt
def microsoft_callback(request):
    code = request.GET.get('code')
    error = request.GET.get('error')
    if error or not code:
        err_msg = f'Lỗi xác thực Microsoft: {error or "Không có mã xác thực"}'
        return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
        
    token_url = "https://login.microsoftonline.com/common/oauth2/v2.0/token"
    data = {
        'client_id': settings.MICROSOFT_CLIENT_ID,
        'scope': 'https://graph.microsoft.com/User.Read email profile openid',
        'code': code,
        'redirect_uri': settings.MICROSOFT_REDIRECT_URI,
        'grant_type': 'authorization_code',
        'client_secret': settings.MICROSOFT_CLIENT_SECRET,
    }
    
    try:
        response = requests.post(token_url, data=data)
        response_data = response.json()
        print("MICROSOFT TOKEN RESPONSE:", response_data)  # Log phản hồi thật của Microsoft
        access_token = response_data.get('access_token')
        if not access_token:
            err_msg = f"Không thể lấy Access Token từ Microsoft. Chi tiết: {response_data.get('error_description', response_data.get('error', 'Không xác định'))}"
            return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
            
        profile_url = "https://graph.microsoft.com/v1.0/me"
        headers = {'Authorization': f'Bearer {access_token}'}
        profile_response = requests.get(profile_url, headers=headers)
        profile = profile_response.json()
        
        email = profile.get('mail') or profile.get('userPrincipalName')
        uid = profile.get('id')
        full_name = profile.get('displayName', 'Microsoft User')
        
        if not email or not uid:
            err_msg = 'Không lấy được email/ID từ tài khoản Microsoft.'
            return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
            
        try:
            user = User.objects.get(oauth_provider='microsoft', oauth_uid=uid)
            login(request, user, backend='accounts.backends.EmailOrUsernameModelBackend')
            return redirect('/html/dashboard.html')
        except User.DoesNotExist:
            if User.objects.filter(email__iexact=email).exists():
                user = User.objects.get(email__iexact=email)
                if not user.oauth_provider:
                    user.oauth_provider = 'microsoft'
                    user.oauth_uid = uid
                    user.save()
                    login(request, user, backend='accounts.backends.EmailOrUsernameModelBackend')
                    return redirect('/html/dashboard.html')
                else:
                    err_msg = 'Email này đã liên kết với một tài khoản xã hội khác.'
                    return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')
            
            request.session['pending_oauth'] = {
                'provider': 'microsoft',
                'uid': uid,
                'email': email,
                'full_name': full_name
            }
            return redirect('/html/login.html?social_register=1')
            
    except Exception as e:
        err_msg = f'Có lỗi xảy ra trong quá trình xác thực Microsoft: {str(e)}'
        return redirect(f'/html/login.html?error_message={urllib.parse.quote(err_msg)}')

@csrf_exempt
def api_oauth_register(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    pending_oauth = request.session.get('pending_oauth')
    if not pending_oauth:
        return JsonResponse({'success': False, 'error': 'Không tìm thấy thông tin xác thực xã hội tạm thời.'}, status=400)
        
    try:
        data = json.loads(request.body)
        username = data.get('username')
        role = data.get('role', 'MEMBER')
        organization = data.get('organization', 'Cá nhân')
        title = data.get('title', '')
        department = data.get('department', '')
    except Exception:
        return JsonResponse({'success': False, 'error': 'Dữ liệu JSON không hợp lệ.'}, status=400)
        
    if not username:
        return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp Tên đăng nhập mong muốn.'}, status=400)
        
    if User.objects.filter(username__iexact=username).exists():
        return JsonResponse({'success': False, 'error': 'Tên đăng nhập này đã tồn tại.'}, status=400)
        
    try:
        user = User.objects.create_user(
            username=username,
            email=pending_oauth['email'],
            password=User.objects.make_random_password(length=18),
            full_name=pending_oauth['full_name'],
            role=role,
            organization=organization or 'Cá nhân',
            title=title,
            department=department,
            oauth_provider=pending_oauth['provider'],
            oauth_uid=pending_oauth['uid']
        )
        
        # Clear session
        del request.session['pending_oauth']
        
        login(request, user, backend='accounts.backends.EmailOrUsernameModelBackend')
        return JsonResponse({
            'success': True,
            'user': {
                'username': user.username,
                'email': user.email,
                'full_name': user.full_name,
                'role': user.role,
                'organization': user.organization
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)


# ==========================================
# DOCUMENT LIFE CYCLE APIs
# ==========================================

@ensure_csrf_cookie
def api_list_documents(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    try:
        documents = Document.objects.filter(user=request.user).order_by('-updated_at')
        docs_data = []
        for doc in documents:
            docs_data.append({
                'id': doc.id,
                'title': doc.title,
                'status': doc.status,
                'created_at': doc.created_at.strftime('%d/%m/%Y %H:%M'),
                'updated_at': doc.updated_at.strftime('%d/%m/%Y %H:%M'),
            })
        return JsonResponse({'success': True, 'documents': docs_data})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)


@ensure_csrf_cookie
def api_get_document(request, doc_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    try:
        doc = Document.objects.filter(user=request.user, id=doc_id).first()
        if not doc:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy tài liệu hoặc bạn không có quyền truy cập.'}, status=404)
            
        return JsonResponse({
            'success': True,
            'document': {
                'id': doc.id,
                'title': doc.title,
                'content': doc.content,
                'status': doc.status,
                'doc_date': doc.doc_date.strftime('%Y-%m-%d'),
                'created_at': doc.created_at.strftime('%d/%m/%Y %H:%M'),
                'updated_at': doc.updated_at.strftime('%d/%m/%Y %H:%M'),
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)


@csrf_exempt
def api_save_document(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        doc_id = data.get('id')
        title = data.get('title', 'Tài liệu không tên').strip()
        content = data.get('content', '')
        status = data.get('status', 'DRAFT')
        doc_date_str = data.get('doc_date')
        
        if not title:
            title = 'Tài liệu không tên'
            
        if status not in ['DRAFT', 'DONE']:
            status = 'DRAFT'
            
        if doc_id:
            doc = Document.objects.filter(user=request.user, id=doc_id).first()
            if not doc:
                return JsonResponse({'success': False, 'error': 'Không tìm thấy tài liệu để cập nhật.'}, status=404)
        else:
            doc = Document(user=request.user)
            
        doc.title = title
        doc.content = content
        doc.status = status
        
        if doc_date_str:
            from django.utils.dateparse import parse_date
            parsed_date = parse_date(doc_date_str)
            if parsed_date:
                doc.doc_date = parsed_date
                
        doc.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Đã lưu tài liệu thành công.',
            'document_id': doc.id,
            'status': doc.status,
            'doc_date': doc.doc_date.strftime('%Y-%m-%d'),
            'updated_at': doc.updated_at.strftime('%d/%m/%Y %H:%M')
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)


@csrf_exempt
def api_delete_document(request, doc_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method not in ['POST', 'DELETE']:
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        doc = Document.objects.filter(user=request.user, id=doc_id).first()
        if not doc:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy tài liệu để xóa.'}, status=404)
            
        doc.delete()
        return JsonResponse({'success': True, 'message': 'Đã xóa tài liệu thành công.'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra: {str(e)}'}, status=500)


@ensure_csrf_cookie
def api_export_docx(request, doc_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    try:
        doc = Document.objects.filter(user=request.user, id=doc_id).first()
        if not doc:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy tài liệu để xuất.'}, status=404)
            
        # Parse HTML to Word using python-docx and BeautifulSoup
        from bs4 import BeautifulSoup
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        import io
        
        docx_doc = DocxDocument()
        
        # Set margins (A4 standard: 2.0 cm Top/Bottom/Right, 3.0 cm Left for binding)
        for section in docx_doc.sections:
            section.top_margin = Inches(0.79)  # 2.0 cm
            section.bottom_margin = Inches(0.79)  # 2.0 cm
            section.left_margin = Inches(1.18)  # 3.0 cm
            section.right_margin = Inches(0.79)  # 2.0 cm
            
        # Configure Normal text style: Times New Roman, 13pt
        style = docx_doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        
        # Parse content
        soup = BeautifulSoup(doc.content, 'html.parser')
        
        for element in soup.contents:
            if element.name is None:  # Direct text node
                text = str(element).strip()
                if text:
                    p = docx_doc.add_paragraph(text)
                    p.paragraph_format.line_spacing = 1.2
                    p.paragraph_format.space_after = Pt(6)
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                heading_text = element.get_text()
                level = int(element.name[1])
                p = docx_doc.add_heading(heading_text, level=level)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                
                # Format heading
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.color.rgb = None  # Default black
                    run.bold = True
                    if level == 1:
                        run.font.size = Pt(16)
                    elif level == 2:
                        run.font.size = Pt(14)
                    else:
                        run.font.size = Pt(13)
            elif element.name == 'p':
                p = docx_doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.space_after = Pt(6)
                
                # Recursive parser for inline formatting
                def parse_inline(node, paragraph):
                    for child in node.children:
                        if child.name is None:
                            run = paragraph.add_run(str(child))
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                        elif child.name in ['strong', 'b']:
                            run = paragraph.add_run(child.get_text())
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                            run.bold = True
                        elif child.name in ['em', 'i']:
                            run = paragraph.add_run(child.get_text())
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                            run.italic = True
                        elif child.name == 'u':
                            run = paragraph.add_run(child.get_text())
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                            run.underline = True
                        else:
                            parse_inline(child, paragraph)
                            
                parse_inline(element, p)
            elif element.name == 'table':
                rows = element.find_all('tr')
                if not rows:
                    continue
                max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                docx_table = docx_doc.add_table(rows=len(rows), cols=max_cols)
                docx_table.style = 'Table Grid'
                
                for r_idx, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for c_idx, cell in enumerate(cells):
                        if c_idx < max_cols:
                            docx_cell = docx_table.cell(r_idx, c_idx)
                            docx_cell.text = cell.get_text().strip()
                            
                            p = docx_cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(2)
                            for run in p.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                                if cell.name == 'th':
                                    run.bold = True
            elif element.name == 'hr':
                docx_doc.add_paragraph().paragraph_format.space_before = Pt(12)
                
        # Save to stream
        file_stream = io.BytesIO()
        docx_doc.save(file_stream)
        file_stream.seek(0)
        
        # Build safe download filename
        filename = doc.title.strip().replace(' ', '_')
        # Remove special characters
        import re
        filename = re.sub(r'[^\w\-_.]', '', filename)
        if not filename:
            filename = f"tai_lieu_{doc.id}"
        filename = f"{filename}.docx"
        
        # URL encode filename for Content-Disposition header (UTF-8 support)
        encoded_filename = urllib.parse.quote(filename.encode('utf-8'))
        
        response = FileResponse(file_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        return response
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Có lỗi xảy ra khi xuất file Word: {str(e)}'}, status=500)


def generate_with_retry(client, model_name, prompt, max_retries=3, delay=5):
    """
    Helper gọi LLM (Ollama Local, OpenAI gpt-4o-mini hoặc Gemini) sinh nội dung
    """
    import os
    import time
    import requests
    import json
    
    # 1. THỬ DÙNG OLLAMA CỤC BỘ (Mô hình Local offline 100% hoàn toàn miễn phí)
    OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://localhost:11434")
    OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "qwen2.5:3b")
    
    try:
        # Kiểm tra nhanh xem cổng dịch vụ Ollama local có đang bật không (timeout 1.5 giây để tránh bị nghẽn)
        chk_resp = requests.get(f"{OLLAMA_URL}/api/tags", timeout=1.5)
        if chk_resp.status_code == 200:
            # Dịch vụ Ollama đang chạy! Kích hoạt sinh văn bản cục bộ
            payload = {
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.2
                }
            }
            # Tăng timeout lên 90 giây vì mô hình lớn chạy trên CPU cần thời gian xử lý
            ollama_resp = requests.post(f"{OLLAMA_URL}/api/generate", json=payload, timeout=90)
            if ollama_resp.status_code == 200:
                res_json = ollama_resp.json()
                # Mock cấu trúc trả về giống response.text của Gemini để tương thích 100% với code xử lý phía dưới
                class OllamaResponseMock:
                    def __init__(self, text):
                        self.text = text
                return OllamaResponseMock(res_json.get("response", ""))
    except Exception as ollama_err:
        # Nếu Ollama chưa được khởi động, log nhẹ và tự động chuyển tiếp sang API đám mây fallback
        print(f"[Ollama Offline] Không thể kết nối tới Ollama tại {OLLAMA_URL} ({ollama_err}). Tự động dùng fallback API...")

    # 2. FALLBACK ĐÁM MÂY (Nếu Ollama chưa được bật)
    OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
    use_openai = OPENAI_API_KEY and not OPENAI_API_KEY.startswith("sk-proj-Your") and OPENAI_API_KEY != "placeholder"
    
    if use_openai:
        from openai import OpenAI
        openai_client = OpenAI(api_key=OPENAI_API_KEY)
        for attempt in range(max_retries):
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.2
                )
                class OpenAIResponseMock:
                    def __init__(self, text):
                        self.text = text
                return OpenAIResponseMock(response.choices[0].message.content)
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"[OpenAI API Error] Sleeping for {delay} seconds and retrying (attempt {attempt + 1}/{max_retries})... Chi tiết: {e}")
                    time.sleep(delay)
                else:
                    raise e
    else:
        # Fallback về Gemini ban đầu
        for attempt in range(max_retries):
            try:
                return client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
            except google_errors.APIError as e:
                if getattr(e, 'code', None) == 429 or "quota" in str(e).lower() or "limit" in str(e).lower():
                    if attempt < max_retries - 1:
                        print(f"[Gemini Quota Exceeded] Sleeping for {delay} seconds and retrying (attempt {attempt + 1}/{max_retries})...")
                        time.sleep(delay)
                    else:
                        raise e
                else:
                    raise e

_e5_model = None

def get_e5_model():
    global _e5_model
    if _e5_model is None:
        from sentence_transformers import SentenceTransformer
        print("⏳ Đang tải mô hình E5 Local 100% offline vào bộ nhớ RAM...")
        _e5_model = SentenceTransformer('intfloat/multilingual-e5-small')
        print("✅ Mô hình E5 Local đã sẵn sàng!")
    return _e5_model


@csrf_exempt
def api_search(request):
    """
    API tìm kiếm thông minh hợp nhất:
    1. Tiếp nhận query và search_mode ("template" hoặc "legal").
    2. Hỗ trợ tiền xử lý nhanh bằng Keyword/Regex để phân loại chính xác, vượt qua LLM khi bị lỗi key.
    3. Luồng "template" (Tìm mẫu văn bản): Tìm kiếm vector và gom nhóm, trả về 5 biểu mẫu liên quan nhất để xem trước toàn văn.
    4. Luồng "legal" (Tra cứu Luật & Nghị định): Tìm kiếm vector trong thư mục files kèm thuật toán cộng điểm Regex, trả về 10 tài liệu liên quan nhất kèm trích dẫn và câu trả lời AI RAG.
    """
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        query = data.get('query', '').strip()
        search_mode = data.get('search_mode', '').strip()  # "template" or "legal"
        
        if not query:
            return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp từ khóa tìm kiếm.'})
            
        import os
        import chromadb
        import re
        from pathlib import Path
        
        # Bảng ánh xạ từ khóa tiếng Việt định danh nhanh
        keyword_mapping = {
            "totrinh": ["tờ trình", "to trinh"],
            "Maudon": ["mẫu đơn", "đơn xin", "đơn đề nghị", "đơn khiếu nại", "đơn tố cáo", "đơn đề xuất", "đơn yêu cầu", "mau don", "đơn "],
            "maubaocao": ["báo cáo", "bao cao", "mẫu báo cáo"],
            "mauphieu": ["phiếu", "mẫu phiếu", "mau phieu", "phiếu thu", "phiếu chi"],
            "thongbao": ["thông báo", "thong bao", "mẫu thông báo"],
            "tokhai": ["tờ khai", "to khai", "mẫu tờ khai"],
            "files": [
                "công văn", "cong van", 
                "nghị định", "nghi dinh", 
                "thông tư", "thong tu", 
                "quyết định", "quyet dinh", 
                "nghị quyết", "nghi quyet", 
                "luật", "luat", 
                "điều khoản", "dieu khoan", 
                "pháp luật", "phap luat", 
                "quy chế", "quy che", 
                "quy định", "quydinh", 
                "chỉ thị", "chi thi", 
                "hướng dẫn", "huong dan"
            ]
        }
        
        # Tiền xử lý phát hiện loại văn bản nhanh
        query_lower = query.lower()
        matched_doc_type = None
        for folder_name, keywords in keyword_mapping.items():
            for kw in keywords:
                if kw in query_lower:
                    matched_doc_type = folder_name
                    break
            if matched_doc_type:
                break
                
        # Tìm thông tin pháp lý bằng regex
        legal_info = {
            "decree_num": None,
            "article_num": None,
            "clause_num": None
        }
        decree_match = re.search(r'(nghị định|nghi dinh|thông tư|thong tu|quyết định|quyet dinh)(\s+số|\s+so)?\s+(\d+([\w\-/]+)?)', query_lower)
        if decree_match:
            legal_info["decree_num"] = decree_match.group(3)
        article_match = re.search(r'(điều|dieu)\s+(\d+)', query_lower)
        if article_match:
            legal_info["article_num"] = article_match.group(2)
        clause_match = re.search(r'(khoản|khoan)\s+(\d+)', query_lower)
        if clause_match:
            legal_info["clause_num"] = clause_match.group(2)
            
        if legal_info["decree_num"] or legal_info["article_num"] or legal_info["clause_num"]:
            matched_doc_type = "files"
            if not search_mode:
                search_mode = "legal"
                
        # Xác định search_mode mặc định nếu chưa chọn
        if not search_mode:
            if matched_doc_type and matched_doc_type != "files":
                search_mode = "template"
            else:
                search_mode = "legal"
                
        # Kết nối database ChromaDB
        BASE_DIR = Path(__file__).resolve().parent.parent
        DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
        
        class NoneEmbeddingFunction:
            def name(self) -> str:
                return "NoneEmbeddingFunction"
            def __call__(self, input):
                return []
                
        chroma_client = chromadb.PersistentClient(path=DB_PATH)
        collection = chroma_client.get_or_create_collection(
            name="knowledge_base",
            embedding_function=NoneEmbeddingFunction()
        )
        
        # Sinh vector embedding của câu truy vấn qua mô hình Local E5-small
        local_model = get_e5_model()
        query_vector = local_model.encode(["query: " + query])[0].tolist()
        
        # ==========================================
        # LUỒNG 1: TÌM MẪU VĂN BẢN (TEMPLATE SEARCH)
        # ==========================================
        if search_mode == 'template':
            # Pass 1: Hệ thống
            query_params = {
                "query_embeddings": [query_vector],
                "n_results": 45
            }
            if matched_doc_type and matched_doc_type != "files":
                query_params["where"] = {"folder": matched_doc_type}
                
            db_results = collection.query(**query_params)
            
            docs_grouped = {}
            if db_results and 'documents' in db_results and db_results['documents'] and len(db_results['documents'][0]) > 0:
                docs = db_results['documents'][0]
                metadatas = db_results['metadatas'][0] if 'metadatas' in db_results else [{}] * len(docs)
                
                for idx, text in enumerate(docs):
                    meta = metadatas[idx] or {}
                    folder = meta.get('folder', '')
                    user_id_meta = meta.get('user_id', '')
                    
                    # Bỏ qua tài liệu pháp luật và tài liệu thuộc về người dùng cụ thể
                    if folder == 'files' or user_id_meta:
                        continue
                        
                    file_name = meta.get('file_name', '')
                    if not file_name:
                        continue
                        
                    if file_name not in docs_grouped:
                        # Tính điểm tượng trưng theo vị trí sắp xếp
                        score = 0.95 - (len(docs_grouped) * 0.03)
                        docs_grouped[file_name] = {
                            "file_name": file_name,
                            "title": clean_filename_to_title(file_name),
                            "folder": folder,
                            "snippet": text[:350] + "...",
                            "score": f"{int(score * 100)}%",
                            "total_chunks": meta.get("total_chunks", 1)
                        }
                        if len(docs_grouped) >= 10:
                            break
                            
            templates_list = list(docs_grouped.values())
            
            # Pass 2: Cá nhân của người dùng hiện tại
            user_templates_list = []
            try:
                user_query_params = {
                    "query_embeddings": [query_vector],
                    "n_results": 20,
                    "where": {"user_id": str(request.user.id)}
                }
                user_db_results = collection.query(**user_query_params)
                
                user_docs_grouped = {}
                if user_db_results and 'documents' in user_db_results and user_db_results['documents'] and len(user_db_results['documents'][0]) > 0:
                    user_docs = user_db_results['documents'][0]
                    user_metadatas = user_db_results['metadatas'][0] if 'metadatas' in user_db_results else [{}] * len(user_docs)
                    
                    for idx, text in enumerate(user_docs):
                        meta = user_metadatas[idx] or {}
                        file_name = meta.get('file_name', '')
                        doc_id = meta.get('doc_id', '')
                        kb_id = meta.get('kb_id', '')
                        
                        if not file_name:
                            continue
                            
                        # Gom nhóm các chunk cùng file cá nhân
                        if file_name not in user_docs_grouped:
                            # Tính điểm
                            score = 0.95 - (len(user_docs_grouped) * 0.04)
                            user_docs_grouped[file_name] = {
                                "file_name": file_name,
                                "title": os.path.splitext(file_name)[0],
                                "doc_id": doc_id,
                                "kb_id": kb_id,
                                "snippet": text[:350] + "...",
                                "score": f"{int(score * 100)}%"
                            }
                            if len(user_docs_grouped) >= 10:
                                break
                    user_templates_list = list(user_docs_grouped.values())
            except Exception as user_search_err:
                print(f"[User Search error] {user_search_err}")
                
            return JsonResponse({
                'success': True,
                'search_mode': 'template',
                'intent': 'TEMPLATE_SEARCH',
                'templates': templates_list,
                'user_templates': user_templates_list
            })
            
        # ==========================================
        # LUỒNG 2: TRA CỨU LUẬT & NGHỊ ĐỊNH (LEGAL QA)
        # ==========================================
        else:
            query_params = {
                "query_embeddings": [query_vector],
                "n_results": 45,
                "where": {"folder": "files"}
            }
            db_results = collection.query(**query_params)
            
            docs_grouped = {}
            scored_chunks = []
            
            if db_results and 'documents' in db_results and db_results['documents'] and len(db_results['documents'][0]) > 0:
                docs = db_results['documents'][0]
                metadatas = db_results['metadatas'][0] if 'metadatas' in db_results else [{}] * len(docs)
                ids = db_results['ids'][0] if 'ids' in db_results else []
                
                for idx, text in enumerate(docs):
                    meta = metadatas[idx] or {}
                    file_name = meta.get('file_name', '').lower()
                    
                    # Tính điểm số lai ghép và cộng điểm thưởng Regex
                    base_score = 0.8 - (idx * 0.005)
                    boost = 0.0
                    
                    # Trừng phạt cực nặng (Demote) các phôi biểu mẫu đính kèm, phụ lục chứa khoảng trống trống để điền
                    # Điều này giúp đẩy các Điều, Khoản quy định pháp lý thực tế (nội dung thực) lên trên cùng!
                    is_annex_or_form = False
                    if "phu luc" in file_name or "mau so" in file_name or "phu-luc" in file_name or "mau-so" in file_name or "bieu-mau" in file_name:
                        is_annex_or_form = True
                    text_stripped = text.strip()
                    if text_stripped.startswith("Mẫu số") or text_stripped.startswith("Phụ lục") or text_stripped.startswith("Phiếu") or text_stripped.startswith("Biên bản") or text_stripped.startswith("Tờ khai"):
                        is_annex_or_form = True
                        
                    if is_annex_or_form:
                        boost -= 0.35 # Trừ điểm mạnh
                    else:
                        # Ưu tiên cộng điểm (Promote) cho các đoạn có cấu trúc quy định Điều/Khoản thực tế
                        if re.search(r'\b(Điều|dieu)\s+\d+', text):
                            boost += 0.15
                        if re.search(r'\b(Khoản|khoan)\s+\d+', text):
                            boost += 0.08
                    
                    # Boost 1: Số hiệu Nghị định/Thông tư
                    if legal_info["decree_num"]:
                        dec_num_str = legal_info["decree_num"].lower()
                        if dec_num_str in file_name:
                            boost += 0.35
                            
                    # Boost 2: Thẩm quyền ký / Thể thức (nếu có từ khóa khớp)
                    if "ký" in query_lower or "thẩm quyền" in query_lower or "thể thức" in query_lower:
                        if "ky" in text.lower() or "tham quyen" in text.lower() or "the thuc" in text.lower():
                            boost += 0.1
                            
                    # Boost 3: Khớp số Điều
                    if legal_info["article_num"]:
                        art_str = f"điều {legal_info['article_num']}"
                        if art_str in text.lower():
                            boost += 0.4
                            
                    # Boost 4: Khớp số Khoản
                    if legal_info["clause_num"]:
                        cls_str = f"khoản {legal_info['clause_num']}"
                        if cls_str in text.lower():
                            boost += 0.15
                            
                    final_score = min(base_score + boost, 0.99)
                    scored_chunks.append({
                        "id": ids[idx] if idx < len(ids) else f"chunk_{idx}",
                        "text": text,
                        "file_name": meta.get('file_name', ''),
                        "folder": meta.get('folder', 'files'),
                        "score": final_score
                    })
                    
                # Sắp xếp theo điểm số sau boost giảm dần
                scored_chunks.sort(key=lambda x: x['score'], reverse=True)
                
                # Gom nhóm Top 10 luật/nghị định độc lập kèm trích dẫn tốt nhất
                for chunk in scored_chunks:
                    f_name = chunk["file_name"]
                    if not f_name:
                        continue
                        
                    if f_name not in docs_grouped:
                        docs_grouped[f_name] = {
                            "file_name": f_name,
                            "title": clean_filename_to_title(f_name),
                            "folder": chunk["folder"],
                            "snippet": chunk["text"][:350] + "...",
                            "score": f"{int(chunk['score'] * 100)}%"
                        }
                        if len(docs_grouped) >= 10:
                            break
                            
            laws_list = list(docs_grouped.values())
            
            # Cấu hình API LLM Fallback để sinh câu trả lời RAG
            GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
            client = None
            if GEMINI_API_KEY:
                client = genai.Client(api_key=GEMINI_API_KEY)
                
            # Tạo ngữ cảnh trả lời từ Top 4 chunks có điểm cao nhất
            top_4_chunks = scored_chunks[:4]
            context = "\n\n".join([f"Tài liệu {c['file_name']}:\n{c['text']}" for c in top_4_chunks])
            
            ans_prompt = f"""
            Bạn là Trợ lý ảo soạn thảo văn bản hành chính Việt Nam cực kỳ thông minh.
            Nhiệm vụ của bạn là trả lời câu hỏi sau của người dùng dựa trên kho tri thức được cung cấp bên dưới:
            
            Yêu cầu:
            1. Trả lời một cách chuyên nghiệp, chính xác, sử dụng thuật ngữ hành chính chuẩn.
            2. Trích dẫn rõ tên văn bản hoặc số hiệu (ví dụ: theo Nghị định 30/2020/NĐ-CP) nếu có trong ngữ cảnh.
            3. Trình bày đẹp mắt sử dụng định dạng Markdown (in đậm để nhấn mạnh các bước hoặc từ khóa quan trọng).
            4. Trả lời trực tiếp, ngắn gọn và tập trung vào câu hỏi.
            
            Ngữ cảnh tri thức:
            {context}
            
            Câu hỏi của người dùng: "{query}"
            """
            
            try:
                ans_response = generate_with_retry(client, 'gemini-2.0-flash', ans_prompt)
                answer_text = ans_response.text.strip()
            except Exception as llm_err:
                print(f"❌ Lỗi sinh câu trả lời RAG: {llm_err}")
                answer_text = "Hệ thống đã định vị được các tài liệu quy định liên quan bên dưới, nhưng dịch vụ tổng hợp câu trả lời AI đang tạm thời gián đoạn. Bạn có thể click 'Xem toàn văn' trực tiếp trên tài liệu để tham khảo."
                
            # Trả về kết quả nguồn dẫn cho RAG tương thích giao diện
            sources = [{
                'name': d['title'],
                'text': d['snippet'],
                'score': d['score'],
                'badge': 'badge-gray',
                'file_name': d['file_name']
            } for d in laws_list[:5]]
            
            return JsonResponse({
                'success': True,
                'search_mode': 'legal',
                'intent': 'LEGAL_QA',
                'answer': answer_text,
                'laws': laws_list,
                'sources': sources
            })
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': f'Lỗi hệ thống tìm kiếm: {str(e)}'}, status=500)


@csrf_exempt
def api_generate_template(request):
    """
    API Truy vấn và Tự động soạn thảo biểu mẫu dựa trên RAG & Prompt chuyên biệt:
    1. Truy vấn biểu mẫu sát nhất trong Vector DB (ChromaDB) để làm context mẫu.
    2. Gọi hàm get_specialized_prompt từ accounts.prompts để lấy Prompt chuyên biệt theo thể loại.
    3. Gửi Prompt và ngữ cảnh mẫu lên Gemini để tự động điền và sinh văn bản cá nhân hóa hoàn chỉnh.
    """
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        option_title = data.get('option_title', '').strip()
        original_query = data.get('original_query', '').strip()
        
        if not option_title:
            return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp loại mẫu cần lấy.'})
            
        import os
        import chromadb
        from pathlib import Path
        from .prompts import get_specialized_prompt
        
        GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
        if not GEMINI_API_KEY:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy GEMINI_API_KEY trong hệ thống.'})
            
        client = genai.Client(api_key=GEMINI_API_KEY)
        
        # 1. Truy vấn biểu mẫu thực tế trong ChromaDB làm ngữ cảnh RAG
        BASE_DIR = Path(__file__).resolve().parent.parent
        DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
        
        class NoneEmbeddingFunction:
            def name(self) -> str:
                return "NoneEmbeddingFunction"
            def __call__(self, input):
                return []
                
        chroma_client = chromadb.PersistentClient(path=DB_PATH)
        collection = chroma_client.get_or_create_collection(
            name="knowledge_base",
            embedding_function=NoneEmbeddingFunction()
        )
        
        # Sinh vector embedding của option_title để truy vấn chính xác qua mô hình Local E5-small
        local_model = get_e5_model()
        query_vector = local_model.encode(["query: " + option_title])[0].tolist()
        
        # Truy vấn mẫu biểu tốt nhất trong ChromaDB (lấy Top 3 chunks)
        db_results = collection.query(
            query_embeddings=[query_vector],
            n_results=3
        )
        
        retrieved_template = ""
        if db_results and 'documents' in db_results and db_results['documents'] and len(db_results['documents'][0]) > 0:
            retrieved_template = "\n\n".join(db_results['documents'][0])
            
        # 2. Lấy Prompt chuyên biệt theo loại văn bản hành chính
        doc_class, prompt = get_specialized_prompt(option_title, original_query, retrieved_template)
        
        # Bổ sung thông tin cá nhân người dùng đang đăng nhập để AI tự động điền mẫu
        user = request.user
        birth_date_str = user.birth_date.strftime('%d/%m/%Y') if user.birth_date else ''
        user_info_context = f"""
        
        [THÔNG TIN THỰC TẾ CỦA NGƯỜI ĐANG SOẠN THẢO VĂN BẢN (DÙNG ĐỂ TỰ ĐỘNG ĐIỀN)]:
        - Họ và tên: {user.full_name or ''}
        - Chức danh/Chức vụ: {user.title or ''}
        - Phòng ban: {user.department or ''}
        - Cơ quan/Tổ chức: {user.organization or ''}
        - Số điện thoại: {user.phone_number or ''}
        - Email: {user.email or ''}
        - Ngày sinh: {birth_date_str}
        
        YÊU CẦU ĐẶC BIỆT: Trong văn bản được tạo, nếu có các vị trí yêu cầu điền thông tin cá nhân (như "Tôi tên là...", "Chức vụ...", "Phòng/Ban...", "Số điện thoại...", "Email...", "Sinh ngày..."), hãy sử dụng các thông tin thực tế bên trên của người soạn thảo để TỰ ĐỘNG ĐIỀN VÀO VĂN BẢN thay vì để trống hoặc để dấu chấm lửng (.....). Việc này giúp cá nhân hóa biểu mẫu ngay lập tức.
        """
        prompt += user_info_context
        
        # 3. Gọi Gemini để thực hiện RAG Generation và tự động soạn thảo
        response = generate_with_retry(client, 'gemini-2.0-flash', prompt)
        
        content_html = response.text.strip()
        # Dọn dẹp markdown blocks nếu LLM vô tình trả về
        if content_html.startswith("```html"):
            content_html = content_html[7:]
        elif content_html.startswith("```"):
            content_html = content_html[3:]
            
        if content_html.endswith("```"):
            content_html = content_html[:-3]
            
        content_html = content_html.strip()
        
        return JsonResponse({
            'success': True,
            'title': option_title,
            'doc_class': doc_class,
            'content': content_html
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Lỗi truy xuất mẫu văn bản RAG: {str(e)}'}, status=500)


def clean_filename_to_title(filename):
    if not filename:
        return "Tài liệu hành chính"
    import os
    import re
    title = os.path.splitext(filename)[0]
    if title.endswith("_chunks"):
        title = title[:-7]
        
    title_upper = title.upper()
    if "ND-CP" in title_upper:
        match = re.search(r'(\d+)-(\d+)-ND-CP', title_upper)
        if match:
            num, year = match.group(1), match.group(2)
            prefix = ""
            if "PHU LUC" in title_upper:
                prefix = "Phụ lục "
            elif "MAU SO" in title_upper:
                prefix = "Mẫu số "
            return f"{prefix}Nghị định {num}/{year}/NĐ-CP"
            
    if "TTLT" in title_upper:
        match = re.search(r'(\d+)-(\d+)-TTLT', title_upper)
        if match:
            num, year = match.group(1), match.group(2)
            return f"Thông tư liên tịch {num}/{year}/TTLT"
            
    if "TT-B" in title_upper or "TT-M" in title_upper or "TT-H" in title_upper:
        match = re.search(r'(\d+)-(\d+)-TT', title_upper)
        if match:
            num, year = match.group(1), match.group(2)
            return f"Thông tư {num}/{year}/TT"
            
    title = title.replace("-", " ").replace("_", " ")
    title = title.strip()
    if title:
        title = title[0].upper() + title[1:]
    return title


@csrf_exempt
def api_get_document_text(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        file_name = data.get('file_name', '').strip()
        if not file_name:
            return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp tên file.'})
            
        import os
        import chromadb
        from pathlib import Path
        
        BASE_DIR = Path(__file__).resolve().parent.parent
        DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
        
        class NoneEmbeddingFunction:
            def name(self) -> str:
                return "NoneEmbeddingFunction"
            def __call__(self, input):
                return []
                
        chroma_client = chromadb.PersistentClient(path=DB_PATH)
        collection = chroma_client.get_or_create_collection(
            name="knowledge_base",
            embedding_function=NoneEmbeddingFunction()
        )
        
        # Lấy toàn bộ các chunk thuộc file này
        db_results = collection.get(
            where={"file_name": file_name},
            limit=150
        )
        
        if not db_results or not db_results.get("documents"):
            return JsonResponse({'success': False, 'error': f'Không tìm thấy nội dung cho file: {file_name}'})
            
        docs = db_results["documents"]
        metadatas = db_results["metadatas"]
        
        chunks_with_idx = []
        for i, text in enumerate(docs):
            meta = metadatas[i] or {}
            try:
                idx = int(meta.get("chunk_index", 0))
            except Exception:
                idx = 0
            chunks_with_idx.append((idx, text))
            
        chunks_with_idx.sort(key=lambda x: x[0])
        full_text = "\n\n".join([item[1] for item in chunks_with_idx])
        
        return JsonResponse({
            'success': True,
            'file_name': file_name,
            'title': clean_filename_to_title(file_name),
            'full_text': full_text
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Lỗi lấy chi tiết văn bản: {str(e)}'}, status=500)


@csrf_exempt
def api_ai_assist(request):
    """
    API Trợ lý ảo AI đắc lực:
    Nhận nội dung hiện tại của khung soạn thảo (HTML) + yêu cầu cải thiện/soạn thảo của người dùng.
    Gọi Gemini 2.0 Flash để tối ưu hóa, sửa lỗi chính tả, hoàn thiện, viết tiếp hoặc điền mẫu thông tin
    trong khi tôn trọng và giữ nguyên cấu trúc thể thức HTML gốc của văn bản hành chính Việt Nam.
    """
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        content = data.get('content', '').strip()
        user_query = data.get('user_query', '').strip()
        doc_title = data.get('title', 'Tài liệu hành chính').strip()
        
        if not user_query:
            return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp yêu cầu cho trợ lý AI.'})
            
        import os
        GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
        if not GEMINI_API_KEY:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy GEMINI_API_KEY trong hệ thống.'})
            
        client = genai.Client(api_key=GEMINI_API_KEY)
        
        user = request.user
        birth_date_str = user.birth_date.strftime('%d/%m/%Y') if user.birth_date else ''
        
        prompt = f"""
        Bạn là một Trợ lý ảo biên tập và tối ưu hóa văn bản hành chính Việt Nam chuyên nghiệp xuất sắc.
        Nhiệm vụ: Chỉnh sửa, hoàn thiện, cải thiện hành văn, sửa lỗi chính tả hoặc tự động điền các thông tin theo [YÊU CẦU NGƯỜI DÙNG] trực tiếp lên [VĂN BẢN HIỆN TẠI].
        
        [THÔNG TIN THỰC TẾ CỦA NGƯỜI ĐANG SOẠN THẢO VĂN BẢN (DÙNG ĐỂ TỰ ĐỘNG ĐIỀN KHI CẦN THIẾT)]:
        - Họ và tên: {user.full_name or ''}
        - Chức danh/Chức vụ: {user.title or ''}
        - Phòng ban: {user.department or ''}
        - Cơ quan/Tổ chức: {user.organization or ''}
        - Số điện thoại: {user.phone_number or ''}
        - Email: {user.email or ''}
        - Ngày sinh: {birth_date_str}
        
        [TÊN TÀI LIỆU]:
        {doc_title}
        
        [VĂN BẢN HIỆN TẠI (ĐỊNH DẠNG HTML)]:
        {content}
        
        [YÊU CẦU NGƯỜI DÙNG]:
        {user_query}
        
        YÊU CẦU BẮT BUỘC:
        1. Giữ nguyên cấu trúc thẻ HTML và thể thức hành chính (Nghị định 30) ban đầu. Chỉ được cập nhật hoặc chỉnh sửa phần nội dung văn bản cho chuyên nghiệp và chính xác theo yêu cầu.
        2. Nếu người dùng yêu cầu điền thông tin (như họ tên, ngày tháng, lý do...) hoặc nếu văn bản có các ô thông tin cá nhân chưa điền, hãy tự động điền bằng thông tin thực tế bên trên của người soạn thảo.
        3. Văn phong hành chính phải chuẩn mực, trang trọng, rõ ràng và đúng chuẩn văn bản nhà nước Việt Nam.
        4. Đầu ra BẮT BUỘC CHỈ là định dạng HTML sạch (chỉ sử dụng các thẻ cơ bản: <h1>, <h2>, <p>, <table>, <tr>, <td>, <strong>, <em>, <u>, <br>). KHÔNG bao gồm <html>, <body> hay CSS phức tạp.
        5. KHÔNG viết lời giải thích dông dài, không bọc mã trong block ```html hay ```. Chỉ trả về duy nhất nội dung mã HTML thô.
        """
        
        response = generate_with_retry(client, 'gemini-2.0-flash', prompt)
        
        content_html = response.text.strip()
        if content_html.startswith("```html"):
            content_html = content_html[7:]
        elif content_html.startswith("```"):
            content_html = content_html[3:]
            
        if content_html.endswith("```"):
            content_html = content_html[:-3]
            
        content_html = content_html.strip()
        
        return JsonResponse({
            'success': True,
            'content': content_html
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Lỗi trợ lý ảo AI biên tập: {str(e)}'}, status=500)


@csrf_exempt
def api_ai_improve(request):
    """
    API Cải thiện văn bản tương tác:
    Phân tích văn bản hiện tại của khung soạn thảo (HTML).
    Gọi Gemini 2.0 Flash để quét và đề xuất tối đa 5-8 điểm cần cải thiện hành văn, sửa lỗi chính tả,
    hoặc chuẩn hóa văn phong theo Nghị định 30.
    Trả về cấu trúc dạng JSON: [{ original, improved, reason }].
    """
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        content = data.get('content', '').strip()
        doc_title = data.get('title', 'Tài liệu hành chính').strip()
        
        if not content or content == '<p><br></p>':
            return JsonResponse({'success': True, 'improvements': []})
            
        import os
        GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
        if not GEMINI_API_KEY:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy GEMINI_API_KEY trong hệ thống.'})
            
        client = genai.Client(api_key=GEMINI_API_KEY)
        
        prompt = f"""
        Bạn là một Trợ lý ảo biên tập văn bản hành chính Việt Nam chuyên nghiệp xuất sắc.
        Nhiệm vụ: Quét qua [VĂN BẢN HIỆN TẠI] và tìm kiếm tối đa 5-8 cụm từ hoặc từ ngữ cần sửa lỗi chính tả, hành văn lủng củng, không trang trọng, hoặc chưa phù hợp với chuẩn phong cách Nghị định 30/2020/NĐ-CP.
        
        [TÊN TÀI LIỆU]:
        {doc_title}
        
        [VĂN BẢN HIỆN TẠI (ĐỊNH DẠNG HTML)]:
        {content}
        
        YÊU CẦU ĐẦU RA BẮT BUỘC:
        1. Chỉ trả về một mảng JSON duy nhất chứa danh sách các đề xuất cải thiện cụ thể. Mỗi đề xuất phải là một đối tượng JSON có đúng 3 trường sau:
           - "original": Cụm từ hoặc từ gốc chính xác xuất hiện trong văn bản (phải trùng khớp từng ký tự bao gồm cả viết hoa/viết thường để hệ thống có thể tìm và thay thế).
           - "improved": Cụm từ cải tiến đề xuất để thay thế.
           - "reason": Giải thích cực kỳ ngắn gọn lý do cải thiện bằng tiếng Việt (Ví dụ: "Sửa lỗi chính tả", "Hành văn trang trọng hơn theo chuẩn công vụ").
        
        2. Mảng JSON phải có cấu trúc như sau:
        [
          {{
            "original": "nghỉ phep",
            "improved": "nghỉ phép",
            "reason": "Sửa lỗi chính tả dấu"
          }}
        ]
        
        3. KHÔNG viết bất kỳ lời giải thích nào khác ngoài mảng JSON này. Cấm bọc mảng trong block ```json hoặc ```. Chỉ trả về chuỗi JSON thô hợp lệ.
        4. Chỉ gợi ý những thay đổi thực sự cần thiết và có độ tin cậy cao để tránh chỉnh sửa lung tung hoặc sai nghĩa gốc. Nếu văn bản đã hoàn hảo, trả về mảng rỗng: [].
        """
        
        response = generate_with_retry(client, 'gemini-2.0-flash', prompt)
        
        response_text = response.text.strip()
        
        # Dọn dẹp markdown blocks nếu LLM vô tình trả về
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        elif response_text.startswith("```"):
            response_text = response_text[3:]
            
        if response_text.endswith("```"):
            response_text = response_text[:-3]
            
        response_text = response_text.strip()
        
        import json as py_json
        try:
            improvements = py_json.loads(response_text)
        except Exception:
            # Fallback if LLM outputted some text around it or formatting issue
            # Try to extract array using regex
            import re
            match = re.search(r'\[\s*\{.*\}\s*\]', response_text, re.DOTALL)
            if match:
                improvements = py_json.loads(match.group(0))
            else:
                improvements = []
                
        return JsonResponse({
            'success': True,
            'improvements': improvements
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Lỗi hệ thống AI cải thiện: {str(e)}'}, status=500)


# =====================================================================
# KNOWLEDGE BASE AND DOCUMENT CRUD APIS WITH USER-LEVEL ISOLATION
# =====================================================================

@csrf_exempt
def api_list_kb(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
    
    kb_list = KnowledgeBase.objects.filter(user=request.user)
    data = []
    for kb in kb_list:
        doc_count = kb.documents.count()
        total_size_bytes = 0
        for doc in kb.documents.all():
            try:
                size_str = doc.file_size
                if "MB" in size_str:
                    total_size_bytes += float(size_str.replace("MB", "").strip()) * 1024 * 1024
                elif "KB" in size_str:
                    total_size_bytes += float(size_str.replace("KB", "").strip()) * 1024
                elif "bytes" in size_str:
                    total_size_bytes += float(size_str.replace("bytes", "").strip())
            except Exception:
                pass
        
        if total_size_bytes >= 1024 * 1024:
            size_formatted = f"{total_size_bytes / (1024*1024):.1f} MB"
        elif total_size_bytes >= 1024:
            size_formatted = f"{total_size_bytes / 1024:.0f} KB"
        else:
            size_formatted = f"{total_size_bytes} bytes"
            
        if doc_count == 0:
            size_formatted = "0 KB"
            
        data.append({
            'id': kb.id,
            'name': kb.name,
            'description': kb.description,
            'doc_count': doc_count,
            'size': size_formatted,
            'created_at': kb.created_at.strftime('%d/%m/%Y'),
        })
    return JsonResponse({'success': True, 'knowledge_bases': data})


@csrf_exempt
def api_create_kb(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        description = data.get('description', '').strip()
        
        if not name:
            return JsonResponse({'success': False, 'error': 'Vui lòng điền tên kho tri thức.'})
            
        kb = KnowledgeBase.objects.create(
            user=request.user,
            name=name,
            description=description
        )
        return JsonResponse({
            'success': True,
            'kb': {
                'id': kb.id,
                'name': kb.name,
                'description': kb.description,
                'doc_count': 0,
                'size': '0 KB',
                'created_at': kb.created_at.strftime('%d/%m/%Y'),
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
def api_delete_kb(request, kb_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'DELETE':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        try:
            kb = KnowledgeBase.objects.get(id=kb_id, user=request.user)
        except KnowledgeBase.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy kho tri thức.'}, status=404)
            
        # Delete from ChromaDB first
        import chromadb
        BASE_DIR = Path(__file__).resolve().parent.parent
        DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
        
        class NoneEmbeddingFunction:
            def name(self) -> str:
                return "NoneEmbeddingFunction"
            def __call__(self, input):
                return []
                
        try:
            chroma_client = chromadb.PersistentClient(path=DB_PATH)
            collection = chroma_client.get_or_create_collection(
                name="knowledge_base",
                embedding_function=NoneEmbeddingFunction()
            )
            collection.delete(where={"kb_id": kb_id})
        except Exception as ce:
            print(f"[ChromaDB Delete Error] {ce}")
            
        # Delete from SQLite
        kb.delete()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
def api_list_kb_documents(request, kb_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    try:
        try:
            kb = KnowledgeBase.objects.get(id=kb_id, user=request.user)
        except KnowledgeBase.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy kho tri thức.'}, status=404)
            
        docs = kb.documents.all()
        data = []
        for doc in docs:
            data.append({
                'id': doc.id,
                'title': doc.title,
                'file_name': doc.file_name,
                'file_type': doc.file_type,
                'file_size': doc.file_size,
                'status': doc.get_status_display(),
                'status_raw': doc.status,
                'created_at': doc.created_at.strftime('%d/%m/%Y'),
            })
        return JsonResponse({'success': True, 'documents': data})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


def extract_pdf_content(pdf_path):
    import pypdf
    reader = pypdf.PdfReader(pdf_path)
    text_list = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_list.append(page_text)
    return "\n\n".join(text_list)


@csrf_exempt
def api_upload_kb_documents(request, kb_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        try:
            kb = KnowledgeBase.objects.get(id=kb_id, user=request.user)
        except KnowledgeBase.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy kho tri thức.'}, status=404)
            
        files = request.FILES.getlist('files')
        if not files:
            return JsonResponse({'success': False, 'error': 'Không nhận được file nào.'})
            
        import chromadb
        from scripts.chunk_pipeline import segment_text_hybrid
        BASE_DIR = Path(__file__).resolve().parent.parent
        sys_path = str(BASE_DIR)
        if sys_path not in sys.path:
            sys.path.append(sys_path)
            
        from scripts.utils_extractor import extract_docx_content, extract_xlsx_content
        
        DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
        
        class NoneEmbeddingFunction:
            def name(self) -> str:
                return "NoneEmbeddingFunction"
            def __call__(self, input):
                return []
                
        chroma_client = chromadb.PersistentClient(path=DB_PATH)
        collection = chroma_client.get_or_create_collection(
            name="knowledge_base",
            embedding_function=NoneEmbeddingFunction()
        )
        
        local_model = get_e5_model()
        uploaded_docs = []
        
        for f in files:
            file_name = f.name
            file_size_bytes = f.size
            
            if file_size_bytes >= 1024 * 1024:
                file_size_str = f"{file_size_bytes / (1024*1024):.1f} MB"
            else:
                file_size_str = f"{file_size_bytes / 1024:.0f} KB"
                
            file_ext = os.path.splitext(file_name)[1].lower()
            
            temp_dir = os.path.join(BASE_DIR, "data", "temp_uploads")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, file_name)
            
            with open(temp_file_path, "wb") as temp_f:
                for chunk in f.chunks():
                    temp_f.write(chunk)
            
            extracted_text = ""
            status = 'active'
            
            try:
                if file_ext == '.pdf':
                    extracted_text = extract_pdf_content(temp_file_path)
                elif file_ext == '.docx':
                    extracted_text = extract_docx_content(temp_file_path)
                elif file_ext == '.xlsx' or file_ext == '.xls':
                    extracted_text = extract_xlsx_content(temp_file_path)
                elif file_ext in ['.txt', '.html', '.htm']:
                    with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as txt_f:
                        raw_text = txt_f.read()
                        if file_ext == '.txt':
                            extracted_text = raw_text
                        else:
                            extracted_text = BeautifulSoup(raw_text, 'html.parser').get_text(separator='\n\n')
                else:
                    with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as txt_f:
                        extracted_text = txt_f.read()
            except Exception as extract_err:
                print(f"[Extraction Error] {file_name}: {extract_err}")
                extracted_text = f"Lỗi trích xuất file {file_name}: {str(extract_err)}"
                status = 'error'
                
            try:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
            except Exception:
                pass
                
            doc = KnowledgeDocument.objects.create(
                kb=kb,
                title=os.path.splitext(file_name)[0],
                file_name=file_name,
                file_type=file_ext[1:].upper() if file_ext else "UNKNOWN",
                file_size=file_size_str,
                status=status,
                content=extracted_text
            )
            
            if status == 'active' and extracted_text.strip():
                chunks = segment_text_hybrid(extracted_text)
                if chunks:
                    passages = ["passage: " + chunk for chunk in chunks]
                    embeddings = local_model.encode(passages, show_progress_bar=False).tolist()
                    
                    ids = [f"kb_doc_{doc.id}_chunk_{i}" for i in range(len(chunks))]
                    metadatas = [{
                        "user_id": str(request.user.id),
                        "kb_id": kb.id,
                        "doc_id": doc.id,
                        "file_name": file_name,
                        "folder": f"kb_{kb.id}"
                    } for _ in range(len(chunks))]
                    
                    collection.add(
                        embeddings=embeddings,
                        documents=chunks,
                        metadatas=metadatas,
                        ids=ids
                    )
            
            uploaded_docs.append({
                'id': doc.id,
                'title': doc.title,
                'file_name': doc.file_name,
                'file_type': doc.file_type,
                'file_size': doc.file_size,
                'status': doc.get_status_display()
            })
            
        return JsonResponse({'success': True, 'uploaded_documents': uploaded_docs})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
def api_delete_kb_document(request, doc_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'DELETE':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        try:
            doc = KnowledgeDocument.objects.get(id=doc_id, kb__user=request.user)
        except KnowledgeDocument.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy tài liệu.'}, status=404)
            
        import chromadb
        BASE_DIR = Path(__file__).resolve().parent.parent
        DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
        
        class NoneEmbeddingFunction:
            def name(self) -> str:
                return "NoneEmbeddingFunction"
            def __call__(self, input):
                return []
                
        try:
            chroma_client = chromadb.PersistentClient(path=DB_PATH)
            collection = chroma_client.get_or_create_collection(
                name="knowledge_base",
                embedding_function=NoneEmbeddingFunction()
            )
            collection.delete(where={"doc_id": doc_id})
        except Exception as ce:
            print(f"[ChromaDB Delete Doc Error] {ce}")
            
        doc.delete()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
def api_get_kb_document_content(request, doc_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    try:
        try:
            doc = KnowledgeDocument.objects.get(id=doc_id, kb__user=request.user)
        except KnowledgeDocument.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy tài liệu.'}, status=404)
            
        from scripts.chunk_pipeline import segment_text_hybrid
        total_chunks = len(segment_text_hybrid(doc.content)) if doc.content.strip() else 0
        
        return JsonResponse({
            'success': True,
            'title': doc.title,
            'file_name': doc.file_name,
            'file_type': doc.file_type,
            'file_size': doc.file_size,
            'content': doc.content,
            'total_chunks': total_chunks,
            'created_at': doc.created_at.strftime('%d/%m/%Y %H:%M')
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
def api_save_editor_to_kb(request, kb_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        try:
            kb = KnowledgeBase.objects.get(id=kb_id, user=request.user)
        except KnowledgeBase.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Không tìm thấy kho tri thức.'}, status=404)
            
        data = json.loads(request.body)
        title = data.get('title', '').strip()
        html_content = data.get('content', '').strip()
        
        if not title:
            return JsonResponse({'success': False, 'error': 'Vui lòng nhập tiêu đề tài liệu.'})
        if not html_content:
            return JsonResponse({'success': False, 'error': 'Nội dung soạn thảo trống.'})
            
        soup = BeautifulSoup(html_content, 'html.parser')
        extracted_text = soup.get_text(separator='\n\n')
        
        size_bytes = len(html_content.encode('utf-8'))
        if size_bytes >= 1024 * 1024:
            file_size_str = f"{size_bytes / (1024*1024):.1f} MB"
        else:
            file_size_str = f"{size_bytes / 1024:.0f} KB"
            
        doc = KnowledgeDocument.objects.create(
            kb=kb,
            title=title,
            file_name=f"{title}.html",
            file_type="HTML",
            file_size=file_size_str,
            status='active',
            content=extracted_text
        )
        
        import chromadb
        from scripts.chunk_pipeline import segment_text_hybrid
        chunks = segment_text_hybrid(extracted_text)
        
        if chunks:
            BASE_DIR = Path(__file__).resolve().parent.parent
            DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
            
            class NoneEmbeddingFunction:
                def name(self) -> str:
                    return "NoneEmbeddingFunction"
                def __call__(self, input):
                    return []
                    
            chroma_client = chromadb.PersistentClient(path=DB_PATH)
            collection = chroma_client.get_or_create_collection(
                name="knowledge_base",
                embedding_function=NoneEmbeddingFunction()
            )
            
            local_model = get_e5_model()
            passages = ["passage: " + chunk for chunk in chunks]
            embeddings = local_model.encode(passages, show_progress_bar=False).tolist()
            
            ids = [f"kb_doc_{doc.id}_chunk_{i}" for i in range(len(chunks))]
            metadatas = [{
                "user_id": str(request.user.id),
                "kb_id": kb.id,
                "doc_id": doc.id,
                "file_name": f"{title}.html",
                "folder": f"kb_{kb.id}"
            } for _ in range(len(chunks))]
            
            collection.add(
                embeddings=embeddings,
                documents=chunks,
                metadatas=metadatas,
                ids=ids
            )
            
        return JsonResponse({
            'success': True,
            'document': {
                'id': doc.id,
                'title': doc.title,
                'file_name': doc.file_name,
                'file_type': doc.file_type,
                'file_size': doc.file_size,
                'status': doc.get_status_display()
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
def api_upload_system_document(request):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Bạn chưa đăng nhập.'}, status=401)
        
    if request.user.role != 'ADMIN':
        return JsonResponse({'success': False, 'error': 'Chỉ có Quản trị viên mới được phép thực hiện chức năng này.'}, status=403)
        
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Phương thức không được hỗ trợ.'}, status=405)
        
    try:
        title = request.POST.get('title', '').strip()
        doc_type = request.POST.get('doc_type', '').strip()
        files = request.FILES.getlist('files')
        
        if not title:
            return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp tiêu đề tài liệu.'})
            
        if not doc_type:
            return JsonResponse({'success': False, 'error': 'Vui lòng cung cấp loại tài liệu.'})
            
        if not files:
            return JsonResponse({'success': False, 'error': 'Vui lòng chọn ít nhất 1 file tài liệu.'})
            
        # Check total file size (limit to 100MB)
        total_size = sum(f.size for f in files)
        if total_size > 100 * 1024 * 1024:
            return JsonResponse({'success': False, 'error': 'Tổng dung lượng tài liệu vượt quá giới hạn 100MB.'})
            
        # Map typed/selected doc_type to folder metadata
        doc_type_clean = doc_type.lower()
        if 'trình' in doc_type_clean:
            matched_folder = 'totrinh'
        elif 'khai' in doc_type_clean:
            matched_folder = 'tokhai'
        elif 'cáo' in doc_type_clean:
            matched_folder = 'maubaocao'
        elif 'đơn' in doc_type_clean:
            matched_folder = 'Maudon'
        elif 'thông báo' in doc_type_clean or 'thong bao' in doc_type_clean:
            matched_folder = 'thongbao'
        elif 'phiếu' in doc_type_clean or 'phieu' in doc_type_clean:
            matched_folder = 'mauphieu'
        else:
            matched_folder = 'files'
            
        import uuid
        import os
        import chromadb
        from scripts.chunk_pipeline import segment_text_hybrid
        from scripts.utils_extractor import extract_docx_content, extract_xlsx_content
        from bs4 import BeautifulSoup
        
        BASE_DIR = Path(__file__).resolve().parent.parent
        DB_PATH = os.path.join(BASE_DIR, "data", "vector_db")
        
        class NoneEmbeddingFunction:
            def name(self) -> str:
                return "NoneEmbeddingFunction"
            def __call__(self, input):
                return []
                
        chroma_client = chromadb.PersistentClient(path=DB_PATH)
        collection = chroma_client.get_or_create_collection(
            name="knowledge_base",
            embedding_function=NoneEmbeddingFunction()
        )
        
        local_model = get_e5_model()
        uploaded_count = 0
        
        for idx, f in enumerate(files):
            orig_file_name = f.name
            file_ext = os.path.splitext(orig_file_name)[1].lower()
            
            # Create a clean file_name that preserves Vietnamese tones
            # Replacing spaces with underscores for clean title recovery
            title_clean = title
            if len(files) > 1:
                title_clean = f"{title}_{idx + 1}"
            
            file_name_clean = title_clean.strip().replace(" ", "_") + file_ext
            
            # Save file to temp
            temp_dir = os.path.join(BASE_DIR, "data", "temp_uploads")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, f"{uuid.uuid4().hex}_{orig_file_name}")
            
            with open(temp_file_path, "wb") as temp_f:
                for chunk in f.chunks():
                    temp_f.write(chunk)
                    
            extracted_text = ""
            try:
                if file_ext == '.pdf':
                    extracted_text = extract_pdf_content(temp_file_path)
                elif file_ext == '.docx':
                    extracted_text = extract_docx_content(temp_file_path)
                elif file_ext == '.xlsx' or file_ext == '.xls':
                    extracted_text = extract_xlsx_content(temp_file_path)
                elif file_ext in ['.txt', '.html', '.htm']:
                    with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as txt_f:
                        raw_text = txt_f.read()
                        if file_ext == '.txt':
                            extracted_text = raw_text
                        else:
                            extracted_text = BeautifulSoup(raw_text, 'html.parser').get_text(separator='\n\n')
                else:
                    with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as txt_f:
                        extracted_text = txt_f.read()
            except Exception as extract_err:
                print(f"[Extraction Error] {orig_file_name}: {extract_err}")
                extracted_text = ""
                
            try:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
            except Exception:
                pass
                
            if extracted_text.strip():
                # Segment text
                chunks = segment_text_hybrid(extracted_text)
                if chunks:
                    # Enforce overwrite if identical file_name exists
                    try:
                        collection.delete(where={"file_name": file_name_clean})
                    except Exception:
                        pass
                        
                    passages = ["passage: " + chunk for chunk in chunks]
                    embeddings = local_model.encode(passages, show_progress_bar=False).tolist()
                    
                    ids = [f"sys_doc_{uuid.uuid4().hex}_chunk_{i}" for i in range(len(chunks))]
                    metadatas = [{
                        "file_name": file_name_clean,
                        "folder": matched_folder,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "user_id": ""  # empty = system knowledge template
                    } for i in range(len(chunks))]
                    
                    collection.add(
                        embeddings=embeddings,
                        documents=chunks,
                        metadatas=metadatas,
                        ids=ids
                    )
                    uploaded_count += 1
                    
        return JsonResponse({'success': True, 'uploaded_count': uploaded_count})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)




